"""Extracteur DOCX via python-docx.

Préserve :
- Hiérarchie des titres (`Heading 1` → `#`, `Heading 2` → `##`, etc.)
- Listes à puces et numérotées
- Tableaux (rendu en Markdown standard)
- Titre du document (métadonnée Core Properties)
"""

from __future__ import annotations

import re
from pathlib import Path

try:
    from docx import Document  # type: ignore[import-untyped]
except ImportError:  # pragma: no cover - dépendance optionnelle absente
    Document = None  # type: ignore[assignment, misc]

from trimtokens.exceptions import MissingDependencyError
from trimtokens.models import ExtractedDocument, ExtractOptions, Section

_HEADING_RE = re.compile(r"^Heading\s+(\d+)$")


def extract(path: Path, options: ExtractOptions) -> ExtractedDocument:
    if Document is None:
        raise MissingDependencyError(
            "Le package 'python-docx' est requis pour extraire les fichiers .docx. "
            "Installez l'extra : pip install 'trimtokens[office]'"
        )
    doc = Document(str(path))

    title = doc.core_properties.title or None

    lines: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = (paragraph.style.name or "") if paragraph.style else ""

        if not text:
            lines.append("")
            continue

        heading_match = _HEADING_RE.match(style_name)
        if heading_match:
            level = max(1, min(6, int(heading_match.group(1))))
            lines.append("#" * level + " " + text)
            continue

        if "Bullet" in style_name or style_name == "List Paragraph":
            lines.append(f"- {text}")
            continue

        if "Number" in style_name:
            lines.append(f"1. {text}")
            continue

        lines.append(text)

    for table in doc.tables:
        rows = list(table.rows)
        if not rows:
            continue
        lines.append("")
        widths = max(len(r.cells) for r in rows)
        for idx, row in enumerate(rows):
            cells = [
                _normalize_cell(row.cells[i].text) if i < len(row.cells) else ""
                for i in range(widths)
            ]
            lines.append("| " + " | ".join(cells) + " |")
            if idx == 0:
                lines.append("| " + " | ".join(["---"] * widths) + " |")
        lines.append("")

    content = "\n".join(lines)

    return ExtractedDocument(
        source_path=path,
        source_type="docx",
        title=title,
        sections=[Section(header="", content=content)],
    )


def _normalize_cell(value: str) -> str:
    return value.strip().replace("|", "\\|").replace("\n", " ")
